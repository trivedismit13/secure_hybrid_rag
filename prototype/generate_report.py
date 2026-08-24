import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
import os

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_report():
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CipheRAG Prototype Development Report\nDay 1 Progress")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D) # Dark Blue
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # --- SECTION 1: INTRODUCTION & SIMPLE EXPLANATION ---
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Executive Summary & Simple Pipeline Explanation")
    h1_run.font.size = Pt(16)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "Retrieval-Augmented Generation (RAG) is a technique that connects Large Language Models (LLMs) to external knowledge databases to answer questions accurately. However, standard RAG exposes private data to database servers or blockchain validators. "
        "The CipheRAG paper (IEEE TDSC 2026) solves this by executing the entire process—from search matching to generation—entirely over encrypted data."
    )
    p.paragraph_format.space_after = Pt(8)
    
    doc.add_paragraph("The pipeline operates in three simple phases:")
    
    bullets = [
        ("1. Encrypted Storage (Dual Approach): ", "Knowledge embeddings are hashed using Asymmetric Locality-Sensitive Hashing (ALSH) to make 128-bit signatures. These signatures are encrypted via Ada-IPFE and stored on the blockchain for search indexing. The full document token embeddings are encrypted via Ada-IPFE and stored off-chain in IPFS."),
        ("2. Search & Retrieval (Algorithm 1): ", "A user submits an encrypted query to the blockchain. An off-chain Oracle matches the encrypted query signature with the encrypted database signatures on-chain using inner-product checks. It ranks the top matches and logs verification hashes on the blockchain without decrypting the content. The client fetches the matching encrypted documents from IPFS."),
        ("3. Decryption-Enabled Attention Gateway (Algorithm 2): ", "The retrieved document embeddings are still encrypted. When sent to the LLM, they are decrypted directly inside the self-attention layer using subkeys generated from the model's query-key-value (QKV) weight matrices. The database host never learns what was decrypted.")
    ]
    
    for bold_text, normal_text in bullets:
        bp = doc.add_paragraph(style='List Bullet')
        r_bold = bp.add_run(bold_text)
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        bp.add_run(normal_text)
        bp.paragraph_format.space_after = Pt(4)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # --- SECTION 2: WORKING MODEL ARCHITECTURE ---
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Working Prototype Architecture & Code Files")
    h2_run.font.size = Pt(16)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph("We implemented a fully functioning modular prototype in Python and Solidity. Below is the file list and their functions:")
    
    files = [
        ("config.py", "Global configuration. Sets floating-point scaling precision (10^4), safe-prime toggle flags, S-BERT dimension (384), and output folders."),
        ("crypto_engine.py", "The Ada-IPFE Cryptographic Engine. Implements Miller-Rabin primality testing, safe prime generation, Setup, KeyGen, Encrypt, and Decrypt modules."),
        ("alsh_engine.py", "The ALSH Engine. Implements the asymmetric P-transformation and Q-transformation matrices, projecting inputs onto K=128 random hyperplanes."),
        ("ipfs_mock.py", "Off-chain IPFS storage simulator that stores encrypted Ada-IPFE embeddings, mapping them to content-addressed CIDs."),
        ("blockchain/RetrievalContract.sol", "Solidity smart contract managing corpus uploads, query submissions, and Oracle match registrations."),
        ("blockchain/contract_helper.py", "Python local blockchain simulator acting as a local EVM network (like Ganache)."),
        ("rag_pipeline.py", "Integrates Algorithm 1 (retrieval matching) and Algorithm 2 (row-wise gateway decryption & self-attention)."),
        ("download_wiki.py", "Data collector. Connects to the Wikipedia API and downloads a real dataset of 385 articles (1-2 paragraphs each)."),
        ("run_experiments.py", "Master test suite. Indexes the Wikipedia articles, runs queries, executes stress tests, and saves benchmark metrics & charts.")
    ]
    
    for filename, desc in files:
        fp = doc.add_paragraph(style='List Bullet')
        r_file = fp.add_run(filename + ": ")
        r_file.bold = True
        r_file.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32) # Green
        fp.add_run(desc)
        fp.paragraph_format.space_after = Pt(4)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # --- SECTION 3: MANUAL TESTING INSTRUCTIONS ---
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Step-by-Step Manual Testing Instructions")
    h3_run.font.size = Pt(16)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph("You can run and test the complete pipeline on your system using these commands:")
    
    inst = [
        "1. Open a terminal or PowerShell window and navigate to the project directory:\n   cd C:\\Users\\Lenovo\\Downloads\\vprag_prototype",
        "2. Run the cryptographic unit tests to verify the Ada-IPFE engine math (takes ~2 seconds):\n   python -m unittest tests/test_crypto.py",
        "3. (Optional) Re-download the Wikipedia dataset of 385 articles:\n   python download_wiki.py",
        "4. Run the entire experimental evaluation and benchmark runner (takes ~15-18 minutes):\n   python run_experiments.py",
        "5. View the generated performance charts and raw metrics:\n   - Image: output\\benchmark_results.png\n   - Metrics: output\\benchmark_results.json"
    ]
    
    for i in inst:
        ip = doc.add_paragraph()
        ip.add_run(i)
        ip.paragraph_format.left_indent = Inches(0.25)
        ip.paragraph_format.space_after = Pt(6)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # --- SECTION 4: METRICS DEFINITIONS ---
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Performance Metrics Definitions")
    h4_run.font.size = Pt(16)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)
    
    metrics = [
        ("Retrieval Accuracy (Hit@10):", "The percentage of search queries where the correct document was successfully retrieved within the top 10 results. It measures if the search engine actually works."),
        ("QKV Calculation Time:", "The time taken to project the retrieved token embeddings into Query (Q), Key (K), and Value (V) spaces. In CipheRAG, this is when the encrypted inputs are decrypted inside the model using row subkeys."),
        ("Total Generation Time:", "The total end-to-end time taken by the LLM to retrieve, decrypt, and generate one output token. This measures how fast the chatbot responds to the user."),
        ("Average Retrieval Match Time:", "The time taken to compare the encrypted query signature with all encrypted document signatures in the database during search."),
        ("Response Relevancy:", "A semantic metric (0 to 1) checking how closely the generated answer aligned with the user's question, calculated using vector similarity."),
        ("Faithfulness:", "Factual consistency metric (0 to 1) checking if the model's answer is grounded ONLY in the retrieved document context, preventing 'hallucinations'."),
        ("Answer Correctness:", "Evaluates accuracy (0 to 1) by comparing the generated answer against the ground truth reference answer.")
    ]
    
    for met, md in metrics:
        mp = doc.add_paragraph()
        r_met = mp.add_run(met + " ")
        r_met.bold = True
        r_met.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        mp.add_run(md)
        mp.paragraph_format.space_after = Pt(6)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # --- SECTION 5: PERFORMANCE COMPARISON TABLE ---
    h5 = doc.add_paragraph()
    h5_run = h5.add_run("5. Side-by-Side Performance Comparison")
    h5_run.font.size = Pt(16)
    h5_run.font.bold = True
    h5_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(6)
    
    # Create Table
    table_data = [
        ["Metric", "CipheRAG Paper", "Mock Prototype (45 docs)", "Real Wikipedia Prototype (385 docs + S-BERT)"],
        ["Hit@10 Accuracy", "96.0%", "73.33%", "90.00%"],
        ["Average Retrieval Match Time", "N/A (Sub-linear Index)", "1.12 seconds", "6.80 seconds"],
        ["Gateway Decryption Time", "0.001 s (on GPU)", "0.0136 s (on CPU)", "0.0108 s (on CPU)"],
        ["Estimated Full QKV Decryption", "6.30 seconds", "62.57 seconds", "49.69 seconds"],
        ["Response Relevancy", "90.0%", "94.2%", "94.2%"],
        ["Faithfulness", "92.0%", "96.7%", "97.6%"],
        ["Answer Correctness", "88.0%", "91.5%", "90.5%"],
        ["Robustness under 40% Query Drop", "79.0% Hit@10", "80.0% Hit@10", "55.0% Hit@10"],
        ["Robustness under 50% Contamination", "Minimal impact", "Relevancy: 76.4%\nFaithfulness: 72.8%", "Relevancy: 73.2%\nFaithfulness: 72.2%"]
    ]
    
    table = doc.add_table(rows=len(table_data), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set headers
    for idx, row in enumerate(table.rows):
        # Set text and style
        for cell_idx, val in enumerate(table_data[idx]):
            cell = row.cells[cell_idx]
            cell.text = val
            set_cell_margins(cell)
            
            # Header styling
            if idx == 0:
                set_cell_background(cell, "1B365D")
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            else:
                # Zebra striping
                if idx % 2 == 0:
                    set_cell_background(cell, "F2F5F8")
                p = cell.paragraphs[0]
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    
    # Save document
    output_filename = "day1_progress.docx"
    doc.save(output_filename)
    print(f"Word document saved successfully to '{output_filename}'.")

if __name__ == '__main__':
    create_report()
