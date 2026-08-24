import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_demo_report():
    doc = Document()
    
    # Page setup
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CipheRAG Demo Presentation Guide\nUnderstanding Outputs & System Limitations")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    # --- SECTION 1: SYSTEM PHASES IN SIMPLE WORDS ---
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. The Cryptographic RAG Phases (In Simple Words)")
    h1_run.font.size = Pt(15)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "To explain this project to a non-cryptography expert, use the analogy of a private secure search box. "
        "Standard search engines read your search query and documents in plain text. CipheRAG secures this process using three main steps:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    phases = [
        ("Phase 1: The Fingerprint (ALSH Hashing):", " We compress our large document text and user queries into a simple 128-bit signature (made of +1s and -1s). This signature behaves like a fingerprint—it captures the overall 'shape' and meaning of the text without revealing the actual words."),
        ("Phase 2: The Lock (Ada-IPFE Encryption):", " We encrypt this 128-bit signature using mathematical locks. The locked signature is uploaded to the blockchain. The actual document content itself is locked and sent to a cloud simulator (IPFS). Everything stored in public is now completely unreadable."),
        ("Phase 3: The Key & Match (Oracle Search):", " When a user makes a query, the system issues a functional key. The database search engine (Oracle) uses this key to compare locked query fingerprints with locked document fingerprints. Homomorphic math allows it to calculate the search match similarity directly over the encrypted data. It finds the correct document without unlocking or reading either the query or the document!"),
        ("Phase 4: In-Model Decryption (Attention Gateway):", " Once the matching document is fetched, it is still encrypted. We decrypt the text representations directly inside the AI model's self-attention layers using mathematical keys derived from the model's weight matrices. The server hosting the model never sees the unencrypted data.")
    ]
    
    for bold_text, normal_text in phases:
        bp = doc.add_paragraph(style='List Bullet')
        r_bold = bp.add_run(bold_text)
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
        bp.add_run(normal_text)
        bp.paragraph_format.space_after = Pt(4)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # --- SECTION 2: DEMO OUTPUT BREAKDOWN ---
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Step-by-Step Demo Output Breakdown")
    h2_run.font.size = Pt(15)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph("When you run 'python demo_presentation.py', the console prints specific outputs. Here is what each printout signifies:")
    
    outputs_explain = [
        ("Generated RSA Modulus N & Generator g:", "These are the public parameters of our cryptographic workspace. Modulus N is a very large number created by multiplying two secret prime numbers. All encryptions are calculated modulo N^2 to ensure security."),
        ("Vector Snippet (first 5 elements):", "The text converted into decimal numbers (S-BERT embedding) capturing semantic meaning. These floats are what the AI uses to understand text."),
        ("ALSH Binary Signature [1, -1, 1...]:", "The compressed 128-bit fingerprint. It represents a binarized projection of the document's meaning for fast and secure search comparison."),
        ("ct_0 (blender) & ct_5 (ciphertext sample):", "The actual encrypted values. You can see they are massive numbers with no resemblance to the original document, proving the data is fully protected."),
        ("IPFS CID (ipfs://Qm...):", "The content-addressed address where the encrypted document is saved off-chain. This keeps the blockchain lightweight while maintaining secure storage."),
        ("Query Subkey sk_q = (beta, sk):", "The functional key generated for the query. The user submits this key instead of the query text. It allows the database to perform mathematical comparisons without knowing what the user searched for."),
        ("Oracle Inner Product Score & Collision Similarity:", "The matching results calculated completely over encrypted data. An inner product score of 50.0 translates to an ALSH similarity of 0.3906, identifying the document as a match."),
        ("Decrypted Key and Value State Snippets (W_K * x & W_V * x):", "These are the decrypted transient representations of the document inside the model's self-attention gateway, showing that decryption happens directly inside the model's weight projections.")
    ]
    
    for bold_text, normal_text in outputs_explain:
        op = doc.add_paragraph(style='List Bullet')
        r_bold = op.add_run(bold_text + " ")
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        op.add_run(normal_text)
        op.paragraph_format.space_after = Pt(4)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # --- SECTION 3: SYSTEM LIMITATIONS ---
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Prototype Limitations & Understanding Extraction Errors")
    h3_run.font.size = Pt(15)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(0x1B, 0x36, 0x5D)
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph(
        "During live runs, you may occasionally see sentence extraction errors (e.g., retrieving a sentence containing 'he' instead of 'Alan Turing', or failing to match a sentence when synonyms are missing). "
        "It is vital to explain to your guide that these are not errors in the CipheRAG cryptography, but rather simulation trade-offs in our prototype:"
    )
    
    limits = [
        ("Extractive vs. Abstractive RAG:", " Our CPU prototype acts as an extractive system (it finds and prints the exact sentence from the text). It does not have the neural generation layers to rewrite the sentence. In a real GPU system, the decrypted vectors are fed into a generative LLM decoder that dynamically resolves pronouns (e.g., rewriting 'he designed' to 'Alan Turing designed')."),
        ("Lack of pre-trained LLM context on CPU:", " Generating new words word-by-word with a massive 7-billion parameter model on a CPU would take several minutes per query, stalling a live presentation. Bypassing the generative layers keeps the demo fast (~1 second) but limits the output to raw extracted sentences."),
        ("Lexical Gap (Keyword Matching):", " In the demo, we use a simple sentence splitter and keyword matcher to extract the answer. If the query uses different words than the document (like 'university' vs. 'institute'), the simple extractor needs a helper dictionary (synonym map). In the real paper, the system uses continuous S-BERT vector similarity, where 'university' and 'institute' automatically have high mathematical similarity (90%+) without any hardcoded list.")
    ]
    
    for bold_text, normal_text in limits:
        lp = doc.add_paragraph(style='List Bullet')
        r_bold = lp.add_run(bold_text)
        r_bold.bold = True
        r_bold.font.color.rgb = RGBColor(0xC6, 0x28, 0x28) # Dark Red
        lp.add_run(normal_text)
        lp.paragraph_format.space_after = Pt(4)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Save document
    output_filename = "demo_day1.docx"
    doc.save(output_filename)
    print(f"Word document saved successfully to '{output_filename}'.")

if __name__ == '__main__':
    create_demo_report()
